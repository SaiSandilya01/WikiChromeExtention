import google.generativeai as genai
from docx import Document
from docx.shared import Inches

def get_gemini_response(prompt_text: str) -> str:
    """
    Sends a prompt to the Gemini 2.0 Flash API and returns the generated text.

    Args:
        prompt_text (str): The text prompt to send to the Gemini API.

    Returns:
        str: The generated text response from the Gemini API.
    """
    # Configure your API key here.
    # IMPORTANT: Replace 'YOUR_API_KEY' with your actual Gemini API key.
    # It's recommended to load your API key from environment variables
    # for security in a real application.
    genai.configure(api_key="AIzaSyBc_3_FzHrIVqx_c1LfTtENKt1ux2YdZL8")

    try:
        # Initialize the Gemini model
        model = genai.GenerativeModel('gemini-2.0-flash')

        # Generate content based on the prompt
        response = model.generate_content(prompt_text)

        # Return the generated text
        if response.candidates and response.candidates[0].content.parts:
            return response.candidates[0].content.parts[0].text
        else:
            return "No content generated or response structure unexpected."
    except Exception as e:
        return f"An error occurred while calling the Gemini API: {e}"

def save_text_to_docx(text_content: str, filename: str = "gemini_response.docx"):
    """
    Saves the given text content into a new Word DOCX file.

    Args:
        text_content (str): The text to be saved in the DOCX file.
        filename (str): The name of the DOCX file to create (e.g., "my_document.docx").
    """
    # Create a new Document
    document = Document()

    # Add a paragraph with the generated text
    document.add_paragraph(text_content)

    # You can add more elements if needed, for example:
    # document.add_heading('Generated Content', level=1)
    # document.add_picture('path/to/image.png', width=Inches(1.25))

    # Save the document
    try:
        document.save(filename)
        print(f"Successfully saved content to '{filename}'")
    except Exception as e:
        print(f"An error occurred while saving the DOCX file: {e}")

# --- Main execution ---
if __name__ == "__main__":
    # 1. Define your prompt for the Gemini API
    user_prompt = "I want to know about nexgen LLC company. give me all references at the end"

    print(f"Sending prompt to Gemini API: '{user_prompt}'")

    # 2. Get the response from Gemini
    gemini_generated_text = get_gemini_response(user_prompt)

    print("\n--- Gemini Response ---")
    print(gemini_generated_text)
    print("-----------------------")

    # 3. Save the Gemini response to a DOCX file
    output_filename = "python_data_analysis_benefits.docx"
    save_text_to_docx(gemini_generated_text, output_filename)

    print(f"\nCheck your current directory for '{output_filename}'")
    